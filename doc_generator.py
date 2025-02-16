# doc_generator.py
import os
import random
from docxtpl import DocxTemplate, InlineImage
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_CELL_VERTICAL_ALIGNMENT
from docx2pdf import convert
from PyPDF2 import PdfMerger
from utils import set_cell_width, chunk_list


def generate_records_doc(record, output_folder):
    template_path = os.path.join(
        "template", "附件1模板", "附件1_自主查核表_首頁模板.docx"
    )
    doc = DocxTemplate(template_path)

    # # 根據 record["district"] 與 record["supervisor_name"] 找到對應資料夾
    # supervisor_folder_path = os.path.join(
    #     os.getcwd(),
    #     "template",
    #     "照片",
    #     record.get("district"),  # e.g. "西區"
    #     "監工",
    #     record.get("supervisor_name"),  # e.g. "朱彥顯"
    # )

    # # 取得資料夾內所有 .jpg 檔案
    # supervisor_image_files = [
    #     f for f in os.listdir(supervisor_folder_path) if f.lower().endswith(".jpg")
    # ]

    # # 如果資料夾內沒有任何 .jpg，直接處理例外或結束
    # if not supervisor_image_files:
    #     print("此資料夾沒有任何 JPG 檔案:", supervisor_folder_path)
    #     return None, None

    # # 從最後兩個檔案中隨機選一個
    # supervisor_name_jpg = random.choice(supervisor_image_files)

    # # 組合成完整路徑
    # supervisor_name_jpg_path = os.path.join(supervisor_folder_path, supervisor_name_jpg)

    # district_folder_path = os.path.join(
    #     os.getcwd(), "template", "照片", record.get("district"), "營業處"  # e.g. "西區"
    # )

    # # 取得資料夾內所有 .jpg 檔案
    # district_image_files = [
    #     f for f in os.listdir(district_folder_path) if f.lower().endswith(".jpg")
    # ]

    # # 如果資料夾內沒有任何 .jpg，直接處理例外或結束
    # if not district_image_files:
    #     print("此資料夾沒有任何 JPG 檔案:", district_folder_path)
    #     return None, None

    # # 從最後兩個檔案中隨機選一個
    # random_jpg = random.choice(district_image_files)

    # # 組合成完整路徑
    # district_path = os.path.join(district_folder_path, random_jpg)
    # 插入到模板中，假設模板裡有 {{ supervisor_name_pic }} 的占位符
    # record["supervisor_name_pic"] = InlineImage(
    #     doc, supervisor_name_jpg_path, width=Cm(6)
    # )
    # record["district_pic"] = InlineImage(doc, district_path, width=Cm(6))
    doc.render(record)
    # 在檔名前加上 context_number
    docx_filename = os.path.join(output_folder, "temp_自主查核表首頁.docx")
    doc.save(docx_filename)
    pdf_path = os.path.join(output_folder, "temp_自主查核表首頁.pdf")
    convert(docx_filename, pdf_path)
    print("Records PDF 已產生：", pdf_path)
    return docx_filename, pdf_path


def generate_pipeline_doc(simulated_data, context_number, output_folder):
    template_path = os.path.join(
        "template", "附件1模板", "附件1_定位資料回饋表_管道模板.docx"
    )
    doc = DocxTemplate(template_path)
    subdoc = doc.new_subdoc()
    num_cols = 7
    table = subdoc.add_table(rows=1, cols=num_cols)
    headers = ["編號", "種類", "座標X", "座標Y", "地盤高程", "埋管深度", "管頂座標z"]
    for i, cell in enumerate(table.rows[0].cells):
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.left_indent = 0
        paragraph.paragraph_format.first_line_indent = 0
        run = paragraph.add_run(headers[i])
        run.font.name = "標楷體"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    header_row = table.rows[0]
    header_row.height = Pt(30)
    header_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    for item in simulated_data:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item["Number"])
        row_cells[1].text = str(item["Type"])
        row_cells[2].text = str(item["Coordinate_X"])
        row_cells[3].text = str(item["Coordinate_Y"])
        row_cells[4].text = str(item["Ground_Elevation"])
        row_cells[5].text = str(item["Pipe_Burial_Depth"])
        row_cells[6].text = str(item["Pipe_Top_Coordinate_Z"])
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.left_indent = 0
                paragraph.paragraph_format.first_line_indent = 0
                for run in paragraph.runs:
                    run.font.name = "標楷體"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")
    tbl = table._element
    tblPr_list = tbl.xpath("./w:tblPr")
    tblPr = tblPr_list[0] if tblPr_list else OxmlElement("w:tblPr")
    if not tblPr_list:
        tbl.insert(0, tblPr)
    tblW_list = tblPr.xpath("./w:tblW")
    tblW = tblW_list[0] if tblW_list else OxmlElement("w:tblW")
    if not tblW_list:
        tblPr.append(tblW)
    tblW.set(qn("w:w"), "10000")
    tblW.set(qn("w:type"), "dxa")
    column_widths = [658, 1756, 1316, 1429, 1094, 1094, 1208]
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, column_widths[idx])
    tbl_borders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement("w:" + border_name)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        tbl_borders.append(border)
    tblPr.append(tbl_borders)
    context = {"table": subdoc, "case_number": context_number}
    doc.render(context)
    word_filename = os.path.join(output_folder, "temp_管線.docx")
    doc.save(word_filename)
    pdf_filename = os.path.join(output_folder, "temp_管線.pdf")
    convert(word_filename, pdf_filename)
    print("管線 PDF 已產生：", pdf_filename)
    return word_filename, pdf_filename


def generate_reserved_doc(reserved_data, context_number, output_folder):
    template_path = os.path.join(
        "template", "附件1模板", "附件1_定位資料回饋表_設施物模板.docx"
    )
    doc = DocxTemplate(template_path)
    subdoc = doc.new_subdoc()
    num_cols = 6
    table = subdoc.add_table(rows=1, cols=num_cols)
    reserved_headers = [
        "編號",
        "種類",
        "座標X",
        "座標Y",
        "地盤高程",
        "座標z",
    ]
    for i, cell in enumerate(table.rows[0].cells):
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.left_indent = 0
        paragraph.paragraph_format.first_line_indent = 0
        run = paragraph.add_run(reserved_headers[i])
        run.font.name = "標楷體"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    header_row = table.rows[0]
    header_row.height = Pt(30)
    header_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    for item in reserved_data:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item["Number"])
        row_cells[1].text = str(item["Type"])
        row_cells[2].text = str(item["Coordinate_X"])
        row_cells[3].text = str(item["Coordinate_Y"])
        row_cells[4].text = str(item["Ground_Elevation"])
        row_cells[5].text = str(item["Pipe_Top_Coordinate_Z"])
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.left_indent = 0
                paragraph.paragraph_format.first_line_indent = 0
                for run in paragraph.runs:
                    run.font.name = "標楷體"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")
    tbl = table._element
    tblPr_list = tbl.xpath("./w:tblPr")
    tblPr = tblPr_list[0] if tblPr_list else OxmlElement("w:tblPr")
    if not tblPr_list:
        tbl.insert(0, tblPr)
    tblW_list = tblPr.xpath("./w:tblW")
    tblW = tblW_list[0] if tblW_list else OxmlElement("w:tblW")
    if not tblW_list:
        tblPr.append(tblW)
    tblW.set(qn("w:w"), "10000")
    tblW.set(qn("w:type"), "dxa")
    column_widths = [658, 1756, 1316, 1429, 1094, 1208]
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, column_widths[idx])
    tbl_borders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement("w:" + border_name)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        tbl_borders.append(border)
    tblPr.append(tbl_borders)
    context = {"table": subdoc, "cast_number": context_number}
    doc.render(context)
    docx_filename = os.path.join(output_folder, "temp_設施物.docx")
    doc.save(docx_filename)
    pdf_filename = os.path.join(output_folder, "temp_設施物.pdf")
    convert(docx_filename, pdf_filename)
    print("設施物 PDF 已產生：", pdf_filename)
    return docx_filename, pdf_filename


def merge_pdf_files(pdf_files, merged_pdf_filename):
    merger = PdfMerger()
    for pdf in pdf_files:
        merger.append(pdf)
    merger.write(merged_pdf_filename)
    merger.close()
    print("PDF 合併完成，最終 PDF 檔案：", merged_pdf_filename)

def generate_image_doc(folder_path, context_number, output_folder):
    template_path = os.path.join("template", "附件4模板", "附件4模板.docx")
    plane_folder = os.path.join(folder_path, "平面圖")
    if not os.path.exists(plane_folder):
        print(f"找不到資料夾：{plane_folder}")
        return None

    valid_exts = (".png", ".jpg", ".jpeg", ".bmp", ".gif")
    image_files = []
    for f in os.listdir(plane_folder):
        if f.lower().endswith(valid_exts):
            name_no_ext, _ = os.path.splitext(f)
            if name_no_ext.isdigit():
                image_files.append(os.path.join(plane_folder, f))
    if not image_files:
        print("在『平面圖』資料夾中找不到符合的照片。")
        return None

    def extract_number(filename):
        base = os.path.basename(filename)
        try:
            return int(os.path.splitext(base)[0])
        except:
            return 0

    image_files.sort(key=extract_number)
    image_groups = [image_files[i : i + 2] for i in range(0, len(image_files), 2)]
    print(f"【平面圖】共找到 {len(image_files)} 張照片，分成 {len(image_groups)} 組。")
    
    if not os.path.exists(template_path):
        print(f"找不到模板檔案：{template_path}")
        return None
    
    temp_pages = []
    for idx, group in enumerate(image_groups, start=1):
        doc = DocxTemplate(template_path)
        subdoc = doc.new_subdoc()
        img_table = subdoc.add_table(rows=2, cols=1)
        
        # 直接設定表格寬度到最大（9020 dxa，大約 15.92 公分）
        # from docx.oxml import OxmlElement
        # from docx.oxml.ns import qn
        tbl = img_table._element
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)
        tblWidth = OxmlElement("w:tblW")
        tblWidth.set(qn("w:w"), "9020")
        tblWidth.set(qn("w:type"), "dxa")
        tblPr.append(tblWidth)
        
        for i in range(2):
            cell = img_table.cell(i, 0)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if i < len(group):
                run = paragraph.add_run()
                run.add_picture(group[i],  height=Cm(7.65))  # 7.65 cm是我自己抓的數字
            else:
                paragraph.add_run("")
        context = {"case_number": context_number, "table": subdoc}
        doc.render(context)
        if idx != len(image_groups):
            doc.add_page_break()
        temp_page = os.path.join(output_folder, f"temp_image_page_{idx}.docx")
        doc.save(temp_page)
        temp_pages.append(temp_page)
        
    merged_doc = Document(temp_pages[0])
    for temp_file in temp_pages[1:]:
        temp_doc = Document(temp_file)
        for element in temp_doc.element.body:
            merged_doc.element.body.append(element)
            
    image_docx_path = os.path.join(output_folder, "temp_平面圖_圖片.docx")
    merged_doc.save(image_docx_path)
    print("【平面圖 - 圖片部分】已儲存:", image_docx_path)
    return image_docx_path


def generate_data_doc(
    simulated_data, reserved_data, context_number, output_folder, max_rows_per_page=50
):
    template_path = os.path.join("template", "附件4模板", "附件4模板.docx")
    if not os.path.exists(template_path):
        print(f"找不到模板檔案：{template_path}")
        return None

    # 使用模板建立文件與子文件
    doc = DocxTemplate(template_path)
    subdoc = doc.new_subdoc()

    # 建立表格（預設 1 行 7 欄，第一行作為 header）
    table = subdoc.add_table(rows=1, cols=7)
    headers = [
        "編號",
        "種類",
        "座標X",
        "座標Y",
        "地盤高程",
        "埋管深度",
        "管頂座標z",
    ]

    # 設定第一個 header 行 (simulated_data 的 header)
    for j, cell in enumerate(table.rows[0].cells):
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.left_indent = 0
        paragraph.paragraph_format.first_line_indent = 0
        run = paragraph.add_run(headers[j])
        run.font.name = "標楷體"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    header_row = table.rows[0]
    header_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    # 插入 simulated_data 資料列
    for row_data in simulated_data:
        row_cells = table.add_row().cells
        row_cells[0].text = str(row_data.get("Number", ""))
        row_cells[1].text = str(row_data.get("Type", ""))
        row_cells[2].text = str(row_data.get("Coordinate_X", ""))
        row_cells[3].text = str(row_data.get("Coordinate_Y", ""))
        row_cells[4].text = str(row_data.get("Ground_Elevation", ""))
        row_cells[5].text = str(row_data.get("Pipe_Burial_Depth", ""))
        row_cells[6].text = str(row_data.get("Pipe_Top_Coordinate_Z", ""))
    if reserved_data:
        headers = [
            "編號",
            "種類",
            "座標X",
            "座標Y",
            "地盤高程",
            "座標z",
            "",
        ]
        # 在 simulated_data 資料列後插入另一個 header 行 (reserved_data 的 header)
        header_cells = table.add_row().cells
        for j, cell in enumerate(header_cells):
            paragraph = cell.paragraphs[0]
            paragraph.clear()  # 清除預設內容以確保格式一致
            run = paragraph.add_run(headers[j])
            run.font.name = "標楷體"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        header_row2 = table.rows[-1]
        header_row2.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

        # 插入 reserved_data 資料列，並移除第6欄資料（埋管深度），將管頂座標z左移
        for row_data in reserved_data:
            row_cells = table.add_row().cells
            row_cells[0].text = str(row_data.get("Number", ""))
            row_cells[1].text = str(row_data.get("Type", ""))
            row_cells[2].text = str(row_data.get("Coordinate_X", ""))
            row_cells[3].text = str(row_data.get("Coordinate_Y", ""))
            row_cells[4].text = str(row_data.get("Ground_Elevation", ""))
            # 將原來應該放在第7欄的「管頂座標z」左移至第6欄
            row_cells[5].text = str(row_data.get("Pipe_Top_Coordinate_Z", ""))
            # 第7欄清空
            row_cells[6].text = ""

    # 設定表格屬性（寬度、邊框等）
    tbl = table._element
    tblPr_list = tbl.xpath("./w:tblPr")
    tblPr = tblPr_list[0] if tblPr_list else OxmlElement("w:tblPr")
    if not tblPr_list:
        tbl.insert(0, tblPr)
    tblW_list = tblPr.xpath("./w:tblW")
    tblW = tblW_list[0] if tblW_list else OxmlElement("w:tblW")
    if not tblW_list:
        tblPr.append(tblW)
    tblW.set(qn("w:w"), "10000")
    tblW.set(qn("w:type"), "dxa")
    tbl_borders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement("w:" + border_name)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        tbl_borders.append(border)
    tblPr.append(tbl_borders)
    column_widths = [658, 1756, 1316, 1429, 1094, 1094, 1208]
    for row in table.rows:
        for idx2, cell in enumerate(row.cells):
            set_cell_width(cell, column_widths[idx2])

    context = {"case_number": context_number, "table": subdoc}
    doc.render(context)
    data_docx_path = os.path.join(output_folder, "temp_平面圖_資料.docx")
    doc.save(data_docx_path)
    print("【平面圖 - 資料部分】已儲存:", data_docx_path)
    return data_docx_path


def merge_docs(doc_paths, output_folder, filename):
    # 以第一個文件作為基礎
    merged_doc = Document(doc_paths[0])
    # 從第二個開始依序加入，每個檔案前加入分頁符
    for doc_path in doc_paths[1:]:
        # merged_doc.add_page_break()
        temp_doc = Document(doc_path)
        for element in temp_doc.element.body:
            merged_doc.element.body.append(element)
    final_docx = os.path.join(output_folder, filename)
    merged_doc.save(final_docx)
    print("【平面圖】最終合併檔案已儲存:", final_docx)
    return final_docx


def merge_pdfs(pdf_paths, output_pdf):
    merger = PdfMerger()
    for pdf in pdf_paths:
        merger.append(pdf)
    merger.write(output_pdf)
    merger.close()
    print(f"【平面圖】PDF 合併完成：{output_pdf}")
