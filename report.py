import os
import re
import json
import tkinter as tk
from tkinter import filedialog

import openpyxl
import pandas as pd
from docxtpl import DocxTemplate, Subdoc, InlineImage
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_CELL_VERTICAL_ALIGNMENT
from docx2pdf import convert
from PyPDF2 import PdfMerger


# -----------------------------------
# 選擇包含 Excel 檔案的資料夾
root = tk.Tk()
root.withdraw()
folder_path = filedialog.askdirectory(title="選擇包含Excel檔案的資料夾")
if not folder_path:
    print("未選擇資料夾，程式結束。")
    exit()

# 從資料夾中找出所有副檔名為 .xlsx 或 .xls 的檔案
excel_files = [
    f for f in os.listdir(folder_path) if f.lower().endswith((".xlsx", ".xls"))
]
if len(excel_files) != 1:
    print("資料夾內必須且僅有一個 Excel 檔案，程式結束。")
    exit()
excel_file_path = os.path.join(folder_path, excel_files[0])
print("選取的 Excel 檔案：", excel_file_path)

# -----------------------------------
# 建立輸出資料夾，輸出資料夾以 case_number 命名
# -----------------------------------
# 這部分會在讀取完 Excel 並處理完資料後再建立 output 資料夾
# 所以此處先不建立，後面根據 df_renamed["case_number"] 再建立

# ============================================
# 第一部份：讀取 Excel，並根據兩種方式整理資料
# ============================================

# ---------------------------
# (1) 利用 pandas 讀取全部工作表資料並進行欄位轉換
# ---------------------------
xls = pd.ExcelFile(excel_file_path)
df = pd.read_excel(xls, sheet_name=xls.sheet_names[0], usecols="A:AA", nrows=2)

# 建立欄位對應的英文名稱
column_mapping = {
    "案號": "case_number",
    "施測日期": "measurement_date",
    "施測人員姓名": "surveyors_name",
    "施測方式": "measurement_method",
    "施測廠商名稱": "survey_company_name",
    "施測廠商電話": "survey_company_phone",
    "技師證號": "technician_license_number",
    "技術士證號": "technician_certificate_number",
    "施測儀器": "survey_equipment",
    "GPS 廠牌型號": "gps_brand_model",
    "經緯儀/全站儀廠牌型號": "total_station_brand_model",
    "潛盾施工廠牌型號": "shield_machine_brand_model",
    "其它廠牌型號": "other_equipment_brand_model",
    "施測點數": "survey_point_count",
    "管線點位": "pipeline_point_count",
    "孔蓋點位": "manhole_point_count",
    "設施物點位": "facility_point_count",
    "參考點位編號": "reference_point_number",
    "參考點位來源": "reference_point_source",
    "原始 E 座標": "original_easting",
    "原始 N 座標": "original_northing",
    "原始 H 正高": "original_height",
    "檢測 E 座標": "measured_easting",
    "檢測 N 座標": "measured_northing",
    "檢測 H 正高": "measured_height",
    "監工名稱": "supervisor_name",
    "區處": "district",
}

# 重新命名 DataFrame 欄位
df_renamed = df.rename(columns=column_mapping)

# 轉換 measurement_date 為 datetime，再拆分成 {year, month, day}
df_renamed["measurement_date"] = pd.to_datetime(
    df_renamed["measurement_date"], errors="coerce"
)
df_renamed["measurement_date"] = df_renamed["measurement_date"].apply(
    lambda x: (
        {"year": x.year, "month": x.month, "day": x.day} if pd.notnull(x) else None
    )
)


# --- 將 measurement_method 轉換成長度為 4 的字串，拆分成四個部分 ---
def transform_measurement_method(x):
    if pd.isnull(x):
        return None
    try:
        s = str(int(x)).zfill(4)  # 轉為 4 位數字字串
    except Exception as e:
        s = "0000"
    return {"part1": s[0], "part2": s[1], "part3": s[2], "part4": s[3]}


df_renamed["measurement_method"] = df_renamed["measurement_method"].apply(
    transform_measurement_method
)
df_renamed["survey_equipment"] = df_renamed["survey_equipment"].apply(
    transform_measurement_method
)

# 建立一個欄位，內容為底線字串（例如22個底線）
df_renamed["underline_22"] = "______________________"

# 將所有 NaN 值預先填入 "empty"
df_renamed = df_renamed.fillna("empty")

# -----------------------------------
# 根據案號建立輸出資料夾（放在 output/下）
# -----------------------------------
output_folder = os.path.join("output", str(df_renamed["case_number"].iloc[0]))
if not os.path.exists(output_folder):
    os.makedirs(output_folder)

# ---------------------------
# (2) 利用 openpyxl 讀取指定範圍並依照 B 欄格式分離資料
# ---------------------------
wb = openpyxl.load_workbook(excel_file_path, data_only=True)
ws = wb.active  # 預設讀取第一個工作表

if not df_renamed.empty:
    context_number = df_renamed["case_number"].iloc[0]
else:
    context_number = None

n_value = df_renamed["survey_point_count"].iloc[0]
if n_value is None:
    print("survey_point_count 欄位沒有數值，程式結束。")
    exit()
try:
    n_value = int(n_value)
except ValueError:
    print("survey_point_count 欄位的數值無法轉換為整數，程式結束。")
    exit()

start_row = 5
end_row = start_row + n_value - 1
data_range = ws[f"A{start_row}:G{end_row}"]

pattern = re.compile(r"^\s*\d+管道點\d+-實測\s*$")

simulated_data = []  # 符合格式的資料
reserved_data = []  # 不符合格式的資料
for row in data_range:
    b_value = row[1].value
    b_str = str(b_value) if b_value is not None else ""
    if not pattern.match(b_str):
        reserved_data.append(
            {
                "Number": row[0].value,
                "Type": row[1].value,
                "Coordinate_X": row[2].value,
                "Coordinate_Y": row[3].value,
                "Ground_Elevation": row[4].value,
                "Pipe_Burial_Depth": row[5].value,
                "Pipe_Top_Coordinate_Z": row[6].value,
            }
        )
        continue

    simulated_data.append(
        {
            "Number": row[0].value,
            "Type": row[1].value,
            "Coordinate_X": (
                round(row[2].value, 4)
                if isinstance(row[2].value, (int, float))
                else row[2].value
            ),
            "Coordinate_Y": (
                round(row[3].value, 4)
                if isinstance(row[3].value, (int, float))
                else row[3].value
            ),
            "Ground_Elevation": (
                round(row[4].value, 3)
                if isinstance(row[4].value, (int, float))
                else row[4].value
            ),
            "Pipe_Burial_Depth": (
                round(row[5].value, 2)
                if isinstance(row[5].value, (int, float))
                else row[5].value
            ),
            "Pipe_Top_Coordinate_Z": (
                round(row[6].value, 4)
                if isinstance(row[6].value, (int, float))
                else row[6].value
            ),
        }
    )

if reserved_data:
    print("以下資料不符合 '數字管道點數字-實測' 格式，將保留起來，不加入主要表格：")
    for item in reserved_data:
        print(item)

# ---------------------------
# (3) （可選）產生 JSON 輸出（此處省略 JSON 產生，可依需求啟用）
# ---------------------------
# 這裡將 JSON 產生部分略過，若需要可參考原程式碼

# ============================================
# 將 records 資料放入首頁模板中 (records_doc)
# ============================================
records_list = df_renamed.to_dict(orient="records")
if records_list:
    record = records_list[0]
else:
    record = {}

records_doc = DocxTemplate("template/附件1模板/附件1_自主查核表_首頁模板.docx")
records_doc.render(record)
records_doc_filename = os.path.join(output_folder, "自主查核表首頁.docx")
records_doc.save(records_doc_filename)
print("Records Word document saved as", records_doc_filename)

# 將首頁 Word 轉成 PDF
records_pdf = os.path.join(output_folder, "自主查核表首頁.pdf")
convert(records_doc_filename, records_pdf)
print("首頁 PDF 已產生：", records_pdf)

# ============================================
# 生成管線模板 (generated_doc)
# ============================================
doc = DocxTemplate("template/附件1模板/附件1_定位資料回饋表_管道模板.docx")
subdoc = doc.new_subdoc()
num_cols = 7
table = subdoc.add_table(rows=1, cols=num_cols)

hdr_cells = table.rows[0].cells
headers = ["編號", "種類", "座標X", "座標Y", "地盤高程", "埋管深度", "管頂座標z"]
for i, text in enumerate(headers):
    paragraph = hdr_cells[i].paragraphs[0]
    paragraph.paragraph_format.left_indent = 0
    paragraph.paragraph_format.first_line_indent = 0
    run = paragraph.add_run(text)
    run.font.name = "標楷體"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

header_row = table.rows[0]
header_row.height = Pt(30)
header_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

for item in simulated_data:
    row_cells = table.add_row().cells
    row_cells[0].text = str(item["Number"])
    row_cells[1].text = str(item["Type"])
    row_cells[2].text = str(item["Coordinate_X"])
    row_cells[3].text = str(item["Coordinate_Y"])
    row_cells[4].text = str(item["Ground_Elevation"])
    row_cells[5].text = str(item["Pipe_Burial_Depth"])
    row_cells[6].text = str(item["Pipe_Top_Coordinate_Z"])
    for cell in row_cells:
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.left_indent = 0
            paragraph.paragraph_format.first_line_indent = 0
            for run in paragraph.runs:
                run.font.name = "標楷體"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")

tbl = table._element
tblPr_list = tbl.xpath("./w:tblPr")
if tblPr_list:
    tblPr = tblPr_list[0]
else:
    tblPr = OxmlElement("w:tblPr")
    tbl.insert(0, tblPr)
tblW_list = tblPr.xpath("./w:tblW")
if tblW_list:
    tblW = tblW_list[0]
else:
    tblW = OxmlElement("w:tblW")
    tblPr.append(tblW)
tblW.set(qn("w:w"), "10000")
tblW.set(qn("w:type"), "dxa")

column_widths = [658, 1756, 1316, 1429, 1094, 1094, 1208]


def set_cell_width(cell, width):
    tc = cell._element
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(width))
    tcW.set(qn("w:type"), "dxa")


for row in table.rows:
    for idx, cell in enumerate(row.cells):
        set_cell_width(cell, column_widths[idx])

tbl_borders = OxmlElement("w:tblBorders")
for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
    border = OxmlElement("w:" + border_name)
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), "8")
    border.set(qn("w:space"), "0")
    border.set(qn("w:color"), "000000")
    tbl_borders.append(border)
tblPr.append(tbl_borders)

context = {
    "table": subdoc,
    "case_number": context_number,
}
doc.render(context)
word_filename = os.path.join(output_folder, "管線.docx")
doc.save(word_filename)
pdf_filename = os.path.join(output_folder, "管線.pdf")
convert(word_filename, pdf_filename)
print("管線 PDF 已產生：", pdf_filename)

# ============================================
# 生成設施物模板 (reserved_doc) (僅當 reserved_data 不為空)
# ============================================
if reserved_data:
    reserved_doc = DocxTemplate(
        "template/附件1模板/附件1_定位資料回饋表_設施物模板.docx"
    )
    reserved_subdoc = reserved_doc.new_subdoc()
    num_cols = 7
    reserved_table = reserved_subdoc.add_table(rows=1, cols=num_cols)

    reserved_headers = [
        "編號",
        "種類",
        "座標X",
        "座標Y",
        "地盤高程",
        "埋管深度",
        "管頂座標z",
    ]
    hdr_cells = reserved_table.rows[0].cells
    for i, text in enumerate(reserved_headers):
        paragraph = hdr_cells[i].paragraphs[0]
        paragraph.paragraph_format.left_indent = 0
        paragraph.paragraph_format.first_line_indent = 0
        run = paragraph.add_run(text)
        run.font.name = "標楷體"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    header_row = reserved_table.rows[0]
    header_row.height = Pt(30)
    header_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    for item in reserved_data:
        row_cells = reserved_table.add_row().cells
        row_cells[0].text = str(item["Number"])
        row_cells[1].text = str(item["Type"])
        row_cells[2].text = str(item["Coordinate_X"])
        row_cells[3].text = str(item["Coordinate_Y"])
        row_cells[4].text = str(item["Ground_Elevation"])
        row_cells[5].text = str(item["Pipe_Burial_Depth"])
        row_cells[6].text = str(item["Pipe_Top_Coordinate_Z"])
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.left_indent = 0
                paragraph.paragraph_format.first_line_indent = 0
                for run in paragraph.runs:
                    run.font.name = "標楷體"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")

    reserved_tbl = reserved_table._element
    reserved_tblPr_list = reserved_tbl.xpath("./w:tblPr")
    if reserved_tblPr_list:
        reserved_tblPr = reserved_tblPr_list[0]
    else:
        reserved_tblPr = OxmlElement("w:tblPr")
        reserved_tbl.insert(0, reserved_tblPr)
    reserved_tblW_list = reserved_tblPr.xpath("./w:tblW")
    if reserved_tblW_list:
        reserved_tblW = reserved_tblW_list[0]
    else:
        reserved_tblW = OxmlElement("w:tblW")
        reserved_tblPr.append(reserved_tblW)
    reserved_tblW.set(qn("w:w"), "10000")
    reserved_tblW.set(qn("w:type"), "dxa")

    for row in reserved_table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, column_widths[idx])

    reserved_tbl_borders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement("w:" + border_name)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        reserved_tbl_borders.append(border)
    reserved_tblPr.append(reserved_tbl_borders)

    reserved_context = {
        "table": reserved_subdoc,
        "cast_number": context_number,
    }
    reserved_doc.render(reserved_context)
    reserved_doc_filename = os.path.join(output_folder, "設施物.docx")
    reserved_doc.save(reserved_doc_filename)
    reserved_pdf_filename = os.path.join(output_folder, "設施物.pdf")
    convert(reserved_doc_filename, reserved_pdf_filename)
    print("設施物 PDF 已產生：", reserved_pdf_filename)

# ============================================
# 合併 PDF 文件
# ============================================
# 首頁與管線 PDF 一定合併，若 reserved_data 不為空則也加入設施物 PDF
merger = PdfMerger()
merger.append(records_pdf)
merger.append(pdf_filename)  # 此為管線 PDF
if reserved_data:
    merger.append(os.path.join(output_folder, "設施物.pdf"))
merged_pdf_filename = os.path.join(output_folder, "附件1.pdf")
merger.write(merged_pdf_filename)
merger.close()
print("PDF 合併完成，最終 PDF 檔案：", merged_pdf_filename)


# ==============================================================================
# 以下為「照片分組」的範例程式碼：
# 需求：利用 tkinter 選一個資料夾，進入該資料夾內的「測量照」子資料夾，
# 將該資料夾中檔名如 pt1、pt2、pt3……的照片每 8 張分為一組，
# 如果照片數量在 8~15 張（即需產生兩頁），或更多時，
# 最後只產生一份包含多頁的 doc 與 pdf 檔案。
# ==============================================================================

# 因為上面的程式已經載入 tkinter、docxtpl、docx2pdf 等模組，
# 這裡只需要另外引入 InlineImage 與圖片尺寸設定所需的模組：


def photo_grouping():
    # ※ 這裡直接使用先前選取的 folder_path
    photo_folder = os.path.join(folder_path, "測量照")
    if not os.path.exists(photo_folder):
        print(f"找不到資料夾：{photo_folder}")
        return

    # 搜尋所有圖片檔案，假設副檔名為 .png, .jpg, .jpeg, .bmp 或 .gif，且檔名以 "pt" 開頭
    image_files = [
        os.path.join(photo_folder, f)
        for f in os.listdir(photo_folder)
        if f.lower().endswith((".png", ".jpg", ".jpeg", ".bmp", ".gif"))
        and f.startswith("pt")
    ]

    if not image_files:
        print("在『測量照』資料夾中找不到符合的照片。")
        return

    # 根據檔名中的數字排序（假設檔名格式為 pt數字.xxx）
    def extract_number(filename):
        base = os.path.basename(filename)
        num_str = "".join(filter(str.isdigit, base))
        try:
            return int(num_str)
        except:
            return 0

    image_files.sort(key=extract_number)

    # 將照片每 8 張分成一組
    groups = [image_files[i : i + 6] for i in range(0, len(image_files), 6)]
    print(f"共找到 {len(image_files)} 張照片，分成 {len(groups)} 組。")

    # 這裡假設模板為單頁模板，位於 "template\附件2模板\附件2模板.docx"
    template_path = (
        r"template\附件2模板\附件2模板.docx"  # 注意使用原始字串避免反斜線問題
    )
    if not os.path.exists(template_path):
        print(f"找不到模板檔案：{template_path}")
        return

    temp_files = []  # 用來存放每一組產生的臨時 docx 檔案
    for idx, group in enumerate(groups, start=1):
        doc = DocxTemplate(template_path)

        # 準備 context 資料，模板中應該有 8 個圖片 placeholder（例如：{{ image_1 }} ~ {{ image_8 }}）以及對應的點位文字（{{ point_1 }} ~ {{ point_8 }})
        context = {"case_number": context_number}
        for i in range(6):
            photo_key = f"image_{i+1}"
            point_key = f"point_{i+1}"
            if i < len(group):
                context[photo_key] = InlineImage(
                    doc, group[i], width=Cm(8.09), height=Cm(5.38)
                )
                context[point_key] = f"編號：pt{i+1}"
            else:
                context[photo_key] = ""
                context[point_key] = ""
        # 渲染模板
        doc.render(context)

        # # 將此組存為臨時檔案
        temp_docx = os.path.join(output_folder, f"temp_photos_{idx}.docx")
        doc.save(temp_docx)
        temp_files.append(temp_docx)

    # 合併所有臨時的 docx 檔案成一份
    merged_doc = Document(temp_files[0])
    # 使用簡單的方式：將後續檔案的 body 內容接續到 merged_doc 中，每組之間插入分頁符
    for temp_file in temp_files[1:]:
        merged_doc.add_page_break()
        temp_doc = Document(temp_file)
        # 直接將 temp_doc 的 XML 內容附加到 merged_doc 的 body
        for element in temp_doc.element.body:
            merged_doc.element.body.append(element)

    merged_docx_filename = os.path.join(output_folder, "測量照.docx")
    merged_doc.save(merged_docx_filename)
    print(f"已儲存合併後的 Docx 檔案: {merged_docx_filename}")

    # 利用 docx2pdf 將合併後的 docx 轉換成 PDF
    merged_pdf_filename = os.path.join(output_folder, "測量照.pdf")
    convert(merged_docx_filename, merged_pdf_filename)
    print(f"已儲存合併後的 PDF 檔案: {merged_pdf_filename}")


# ==============================================================================
# 以下為第三個模板的照片處理流程：
# 需求：進入資料夾中的「讀數照」子資料夾，
# 檔名規則由 "pt" 改為 "app_pt"，
# 模板路徑為第三個模板（例如：template\附件3模板\附件3模板.docx），
# 最後只產生一份包含多頁的 doc 與 pdf 檔案。
# ==============================================================================


def photo_grouping_app():
    # 指定「讀數照」子資料夾
    photo_folder = os.path.join(folder_path, "讀數照")
    if not os.path.exists(photo_folder):
        print(f"找不到資料夾：{photo_folder}")
        return

    # 搜尋所有圖片檔案，假設副檔名為 .png, .jpg, .jpeg, .bmp 或 .gif，且檔名以 "app_pt" 開頭
    image_files = [
        os.path.join(photo_folder, f)
        for f in os.listdir(photo_folder)
        if f.lower().endswith((".png", ".jpg", ".jpeg", ".bmp", ".gif"))
        and f.startswith("app_pt")
    ]

    if not image_files:
        print("在『讀數照』資料夾中找不到符合的照片。")
        return

    # 根據檔名中的數字排序（假設檔名格式為 app_pt數字.xxx）
    def extract_number(filename):
        base = os.path.basename(filename)
        num_str = "".join(filter(str.isdigit, base))
        try:
            return int(num_str)
        except:
            return 0

    image_files.sort(key=extract_number)

    # 將照片每 6 張分成一組（例如照片數量在 8~15 張時，會產生 2 頁）
    groups = [image_files[i : i + 6] for i in range(0, len(image_files), 6)]
    print(f"【讀數照】共找到 {len(image_files)} 張照片，分成 {len(groups)} 組。")

    # 第三個模板的路徑，假設為 "template\附件3模板\附件3模板.docx"
    template_path = r"template\附件3模板\附件3模板.docx"  # 使用原始字串避免反斜線問題
    if not os.path.exists(template_path):
        print(f"找不到模板檔案：{template_path}")
        return

    temp_files = []  # 用來存放每一組產生的臨時 docx 檔案
    for idx, group in enumerate(groups, start=1):
        doc = DocxTemplate(template_path)

        # 準備 context 資料，模板中應有 6 個圖片 placeholder（例如：{{ image_1 }} ~ {{ image_6 }}）
        # 以及對應的點位文字（例如：{{ point_1 }} ~ {{ point_6 }}）
        context = {"case_number": context_number}
        for i in range(6):
            photo_key = f"image_{i+1}"
            point_key = f"point_{i+1}"
            if i < len(group):
                # 指定圖片尺寸（這裡的尺寸請根據第三個模板中格子的實際尺寸設定）
                context[photo_key] = InlineImage(
                    doc, group[i], width=Cm(8.09), height=Cm(5.38)
                )
                context[point_key] = f"編號：app_pt{i+1}"
            else:
                context[photo_key] = ""
                context[point_key] = ""
        doc.render(context)

        # 將此組存為臨時檔案
        temp_docx = os.path.join(output_folder, f"temp_app_photos_{idx}.docx")
        doc.save(temp_docx)
        temp_files.append(temp_docx)

    # 合併所有臨時的 docx 檔案成一份
    merged_doc = Document(temp_files[0])
    for temp_file in temp_files[1:]:
        merged_doc.add_page_break()
        temp_doc = Document(temp_file)
        for element in temp_doc.element.body:
            merged_doc.element.body.append(element)

    merged_docx_filename = os.path.join(output_folder, "讀數照.docx")
    merged_doc.save(merged_docx_filename)
    print(f"【讀數照】已儲存合併後的 Docx 檔案: {merged_docx_filename}")

    # 利用 docx2pdf 將合併後的 docx 轉換成 PDF
    merged_pdf_filename = os.path.join(output_folder, "讀數照.pdf")
    convert(merged_docx_filename, merged_pdf_filename)
    print(f"【讀數照】已儲存合併後的 PDF 檔案: {merged_pdf_filename}")


# ==============================================================================
# 以下為最後一個模板的處理程式碼：
# 需求：讀取資料夾「平面圖」中檔名僅為數字的照片，
# 每一頁的 Docx 僅有一個 1x2 的表格，並將兩張照片分別插入表格左、右兩個儲存格中，
# 最後產生一份包含多頁的 Docx 與 PDF 檔案。
# ==============================================================================

# =============================================================================
# 假設下列全域變數由前面流程設定（例如 Excel 讀取後）
# 請根據你的實際程式碼來定義：
#   folder_path: 使用者選取的資料夾路徑
#   context_number: 例如從 Excel 中讀取到的案號
#   simulated_data: 從 Excel 處理出來的管線資料列表
#   reserved_data: 從 Excel 處理出來的不符合格式的資料列表
#   output_folder: 輸出資料夾（以案號命名）－例如：
#       output_folder = os.path.join("output", context_number)
# 另外，本範例將 simulated_data 與 reserved_data 合併為 combined_data
# =============================================================================

# 以下變數請替換成你原有流程產生的變數
# 例如：
# folder_path = "C:/Users/YourUser/SomeFolder"
# context_number = "案號123"
# simulated_data = [...]      # 例如：[{ "Number": ..., "Type": ..., ...}, ...]
# reserved_data = [...]       # 同上
# output_folder = os.path.join("output", context_number)
# -----------------------------------------------------------------------------
# 為方便示範，這裡假設：


# -----------------------------------------------------------------------------


def chunk_list(data_list, chunk_size):
    """將 data_list 分割成每塊不超過 chunk_size 的子列表"""
    return [data_list[i : i + chunk_size] for i in range(0, len(data_list), chunk_size)]


# =============================================================================
# ① 產生圖片專用的 Doc
# =============================================================================
def generate_image_doc():
    template_path = r"template\附件4模板\附件4模板.docx"  # 請根據實際路徑調整
    if not os.path.exists(template_path):
        print(f"找不到模板檔案：{template_path}")
        return None

    # 讀取「平面圖」子資料夾中的圖片（檔名僅為數字）
    plane_folder = os.path.join(folder_path, "平面圖")
    if not os.path.exists(plane_folder):
        print(f"找不到資料夾：{plane_folder}")
        return None

    image_files = []
    for f in os.listdir(plane_folder):
        if f.lower().endswith((".png", ".jpg", ".jpeg", ".bmp", ".gif")):
            name_no_ext, _ = os.path.splitext(f)
            if name_no_ext.isdigit():
                image_files.append(os.path.join(plane_folder, f))
    if not image_files:
        print("在『平面圖』資料夾中找不到符合的照片。")
        return None

    # 根據檔名數字排序
    def extract_number(filename):
        base = os.path.basename(filename)
        num_str = os.path.splitext(base)[0]
        try:
            return int(num_str)
        except:
            return 0

    image_files.sort(key=extract_number)
    # 每2張為一組
    image_groups = [image_files[i : i + 2] for i in range(0, len(image_files), 2)]
    print(f"【平面圖】共找到 {len(image_files)} 張照片，分成 {len(image_groups)} 組。")

    temp_pages = []  # 用來存放每一頁的暫存檔路徑
    for idx, group in enumerate(image_groups):
        page_tpl = DocxTemplate(template_path)
        subdoc = page_tpl.new_subdoc()
        # 建立一個 1×2 的表格，用於放圖片
        img_table = subdoc.add_table(rows=2, cols=1)
        img_table.style = "Table Grid"
        for i in range(2):
            cell = img_table.cell(0, i)
            paragraph = cell.paragraphs[0]
            if i < len(group):
                run = paragraph.add_run()
                # 請根據實際格子尺寸調整（此例寬 12 公分，高 10 公分）
                run.add_picture(group[i], width=Cm(12), height=Cm(10))
            else:
                paragraph.add_run("")
        # 只填入圖片區域，資料區域留空
        context = {"case_number": context_number, "table": subdoc}
        page_tpl.render(context)
        temp_page = os.path.join(output_folder, f"temp_image_page_{idx+1}.docx")
        page_tpl.save(temp_page)
        temp_pages.append(temp_page)

    # 合併所有圖片頁
    merged_img_doc = Document(temp_pages[0])
    for temp_file in temp_pages[1:]:
        merged_img_doc.add_page_break()
        temp_doc = Document(temp_file)
        for element in temp_doc.element.body:
            merged_img_doc.element.body.append(element)
    image_docx_path = os.path.join(output_folder, "平面圖_圖片.docx")
    merged_img_doc.save(image_docx_path)
    print(f"【平面圖 - 圖片部分】已儲存: {image_docx_path}")
    return image_docx_path


# =============================================================================
# ② 產生資料專用的 Doc（管線／設施物資料）
# =============================================================================
def generate_data_doc(max_rows_per_page=10):
    template_path = r"template\附件4模板\附件4模板.docx"  # 同上
    if not os.path.exists(template_path):
        print(f"找不到模板檔案：{template_path}")
        return None

    # 將 combined_data 依 max_rows_per_page 分割成資料區塊
    combined_data = simulated_data + reserved_data
    data_chunks = chunk_list(combined_data, max_rows_per_page)
    print(
        f"【平面圖 - 資料部分】共分成 {len(data_chunks)} 個資料區塊（每頁最多 {max_rows_per_page} 行）。"
    )

    temp_pages = []
    for idx, chunk in enumerate(data_chunks):
        page_tpl = DocxTemplate(template_path)
        subdoc = page_tpl.new_subdoc()
        # 建立資料表格，假設有 7 欄（依你實際資料調整）
        data_table = subdoc.add_table(rows=1, cols=7)
        data_table.style = "Table Grid"
        headers = [
            "編號",
            "種類",
            "座標X",
            "座標Y",
            "地盤高程",
            "埋管深度",
            "管頂座標z",
        ]
        hdr_cells = data_table.rows[0].cells
        # 依據管線模板範例，對每個 header cell 進行格式化
        for j, header in enumerate(headers):
            paragraph = hdr_cells[j].paragraphs[0]
            paragraph.paragraph_format.left_indent = 0
            paragraph.paragraph_format.first_line_indent = 0
            run = paragraph.add_run(header)
            run.font.name = "標楷體"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            hdr_cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        # 設定 header row 高度
        header_row = data_table.rows[0]
        header_row.height = Pt(30)
        header_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        for row_data in chunk:
            row_cells = data_table.add_row().cells
            row_cells[0].text = str(row_data.get("Number", ""))
            row_cells[1].text = str(row_data.get("Type", ""))
            row_cells[2].text = str(row_data.get("Coordinate_X", ""))
            row_cells[3].text = str(row_data.get("Coordinate_Y", ""))
            row_cells[4].text = str(row_data.get("Ground_Elevation", ""))
            row_cells[5].text = str(row_data.get("Pipe_Burial_Depth", ""))
            row_cells[6].text = str(row_data.get("Pipe_Top_Coordinate_Z", ""))
        # 只填入資料區塊，圖片區域留空

        tbl = data_table._element
        tblPr_list = tbl.xpath("./w:tblPr")
        if tblPr_list:
            tblPr = tblPr_list[0]
        else:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)
        tblW_list = tblPr.xpath("./w:tblW")
        if tblW_list:
            tblW = tblW_list[0]
        else:
            tblW = OxmlElement("w:tblW")
            tblPr.append(tblW)
        tblW.set(qn("w:w"), "10000")
        tblW.set(qn("w:type"), "dxa")

        # 使用與其他模板相同的欄寬設定
        column_widths = [658, 1756, 1316, 1429, 1094, 1094, 1208]
        for row in data_table.rows:
            for idx, cell in enumerate(row.cells):
                set_cell_width(cell, column_widths[idx])

        context = {"case_number": context_number, "table": subdoc}
        page_tpl.render(context)
        temp_page = os.path.join(output_folder, f"temp_data_page_{idx+1}.docx")
        page_tpl.save(temp_page)
        temp_pages.append(temp_page)

    merged_data_doc = Document(temp_pages[0])
    for temp_file in temp_pages[1:]:
        merged_data_doc.add_page_break()
        temp_doc = Document(temp_file)
        for element in temp_doc.element.body:
            merged_data_doc.element.body.append(element)
    data_docx_path = os.path.join(output_folder, "平面圖_資料.docx")
    merged_data_doc.save(data_docx_path)
    print(f"【平面圖 - 資料部分】已儲存: {data_docx_path}")
    return data_docx_path


# =============================================================================
# ③ 合併圖片 Doc 與 資料 Doc
# =============================================================================
def merge_docs(doc1_path, doc2_path):
    merged_doc = Document(doc1_path)
    merged_doc.add_page_break()
    temp_doc = Document(doc2_path)
    for element in temp_doc.element.body:
        merged_doc.element.body.append(element)
    final_docx = os.path.join(output_folder, "平面圖_final.docx")
    merged_doc.save(final_docx)
    print(f"【平面圖】最終合併檔案已儲存: {final_docx}")
    return final_docx


if __name__ == "__main__":
    # 這邊會先執行上方的 Excel/Word/PDF 產生流程，
    # 完成後依序執行照片分組的流程，產出第二個模板（測量照）的文件，
    # 以及第三個模板（讀數照）的文件。
    print("========== Excel 相關文件產生完成 ==========")
    print("========== 現在開始『測量照』照片分組處理 ==========")
    photo_grouping()
    print("========== 『讀數照』照片分組處理 ==========")
    photo_grouping_app()
    print("========== 現在開始產生平面圖 - 圖片部分 ==========")
    image_docx = generate_image_doc()
    print("========== 現在開始產生平面圖 - 資料部分 ==========")
    data_docx = generate_data_doc(max_rows_per_page=10)
    if image_docx and data_docx:
        final_docx = merge_docs(image_docx, data_docx)
        final_pdf = os.path.join(output_folder, "平面圖_final.pdf")
        convert(final_docx, final_pdf)
        print(f"【平面圖】最終 PDF 已儲存: {final_pdf}")
